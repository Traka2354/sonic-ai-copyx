from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import (Font, Alignment, PatternFill, Border, Side,
                              GradientFill)
from openpyxl.utils import get_column_letter
import copy


# ─────────────────────────────────────────────
#  HELPERS
# ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    """Set table-cell background colour (Word)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        val = kwargs.get(side, {})
        if val:
            tag = OxmlElement(f'w:{side}')
            for k, v in val.items():
                tag.set(qn(f'w:{k}'), v)
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def box_border(cell, size='12', color='000000'):
    b = {'val': 'single', 'sz': size, 'space': '0', 'color': color}
    set_cell_border(cell, top=b, left=b, bottom=b, right=b)


def add_run(para, text, bold=False, italic=False, size=9,
            color=None, underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return run


def cell_para(cell, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Return first paragraph of a cell, cleared, with alignment set."""
    p = cell.paragraphs[0]
    p.alignment = alignment
    return p


def add_para(cell, text='', bold=False, italic=False, size=9,
             color=None, underline=False,
             alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0):
    """Add a new paragraph to a cell (not clearing existing ones)."""
    para = cell.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    add_run(para, text, bold=bold, italic=italic,
            size=size, color=color, underline=underline)
    return para


# ─────────────────────────────────────────────
#  WORD DOCUMENT
# ─────────────────────────────────────────────
def create_word():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)

    # ── HEADER TABLE  (title left | form fields right) ──────────────────
    hdr = doc.add_table(rows=1, cols=2)
    hdr.style = 'Table Grid'
    hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr.columns[0].width = Inches(3.2)
    hdr.columns[1].width = Inches(3.8)

    # Left – big title
    lc = hdr.cell(0, 0)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(lc, top={}, left={}, bottom={}, right={})  # no border left cell

    p = cell_para(lc)
    add_run(p, 'ONBOARD\nEXPERIENCES', bold=True, size=28)
    p2 = lc.add_paragraph()
    add_run(p2, 'CORFU, GREECE', bold=True, size=11)
    p3 = lc.add_paragraph()
    add_run(p3, 'APPROACH OUR TEAM FOR ASSISTANCE', bold=True, size=8)

    # Right – form fields
    rc = hdr.cell(0, 1)
    set_cell_border(rc, top={}, left={}, bottom={}, right={})

    for label in ('GUEST NAME:', 'ROOM NUMBER:', "CREWMEMBER'S NAME & MAPS ID:"):
        ft = rc.add_table(rows=1, cols=1)
        ft.style = 'Table Grid'
        fc = ft.cell(0, 0)
        box_border(fc, size='6')
        fp = cell_para(fc)
        add_run(fp, label, bold=True, size=8)
        fc.add_paragraph()  # blank line inside box
        rc.add_paragraph()  # spacing between boxes

    doc.add_paragraph()  # spacer

    # ── MAIN GRID  (3 columns) ───────────────────────────────────────────
    # Row 1: Shore Excursion | Vibe Beach Club
    # Row 2: Airport Transfer | Thermal Spa
    # Row 3: (empty) | Train Station Transfer | Laundry Promo

    main = doc.add_table(rows=3, cols=3)
    main.style = 'Table Grid'
    main.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_w = [Inches(2.2), Inches(0.05), Inches(4.75)]  # approx widths
    for i, w in enumerate(col_w):
        for row in main.rows:
            row.cells[i].width = w

    # ── SHORE EXCURSION (rows 0-1, col 0) ───────────────────────────────
    shore = main.cell(0, 0).merge(main.cell(1, 0))
    box_border(shore)
    set_cell_bg(shore, 'FFFFFF')
    shore.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    p = cell_para(shore, WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, 'SHORE EXCURSION', bold=True, size=9, color='FFFFFF')
    set_cell_bg(shore, '333333')

    # Because we already set bg above, we re-do: use nested table approach
    # ---- simpler: just colour the header paragraph via shading the cell
    # Reset and rebuild with a header row approach via nested table
    shore = main.cell(0, 0).merge(main.cell(1, 0))
    shore._tc.clear_content()  # wipe
    set_cell_bg(shore, 'FFFFFF')
    shore.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    box_border(shore)

    # Header bar via shaded paragraph trick — use a nested 1-col table
    nt = shore.add_table(rows=1, cols=1)
    nt.style = 'Table Grid'
    nhc = nt.cell(0, 0)
    set_cell_bg(nhc, '333333')
    box_border(nhc, color='333333')
    np_ = cell_para(nhc, WD_ALIGN_PARAGRAPH.CENTER)
    add_run(np_, 'SHORE EXCURSION', bold=True, size=9, color='FFFFFF')

    def shore_entry(cell, title, rows):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, title, bold=False, underline=True, size=9)
        for line in rows:
            pp = cell.add_paragraph()
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(pp, line, bold='$' in line, size=8)

    shore_entry(shore, 'Kotor Tours', [
        'BOKA BAY CATAMARAN SAIL & BEACH',
        '$199.99Adult | $149.99 Child',
        'HISTORIC KOTOR & BUDVA',
        '$99.99Adult | $79.99 Child',
    ])
    shore_entry(shore, 'Split Tours', [
        'TROGIR & BEACH',
        '$89.99Adult | $49.99 Child',
        'SPLIT OPEN-AIR SIGHTSEEING\n& WALKING TOUR',
        '$69.99Adult | $49.99 Child',
    ])
    fp = shore.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(fp, 'Visit Shore Excursions Desk\nDeck 7, Mid', italic=True, size=8)

    # ── middle spacer column ─────────────────────────────────────────────
    for r in range(3):
        c = main.cell(r, 1)
        set_cell_border(c, top={}, left={}, bottom={}, right={})

    # ── VIBE BEACH CLUB (row 0, col 2) ──────────────────────────────────
    vibe = main.cell(0, 2)
    box_border(vibe)
    set_cell_bg(vibe, 'FFFFFF')
    vibe.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    nvt = vibe.add_table(rows=1, cols=1)
    nvt.style = 'Table Grid'
    nvhc = nvt.cell(0, 0)
    set_cell_bg(nvhc, '333333')
    box_border(nvhc, color='333333')
    nvp = cell_para(nvhc, WD_ALIGN_PARAGRAPH.CENTER)
    add_run(nvp, 'VIBE BEACH CLUB', bold=True, size=9, color='FFFFFF')

    for line in [
        ('Escape to luxury at sea.', False, False, 8),
        ('Unwind in style with an exclusive, adults-only retreat\nwith ocean views & enjoy premium touches designed\njust for you.', False, False, 8),
        ('Limited access available.', False, False, 8),
        ('Reserve early for the ultimate experience.', False, False, 8),
        ('Vibe Cruise Pass - $209/person (4 days)', True, False, 8),
        ('Vibe Cabanas - $499/2pax (4 days)', True, False, 8),
        ('', False, False, 8),
        ('Visit Guest Service Desk, Deck 7 Midship', False, True, 8),
    ]:
        vp = vibe.add_paragraph()
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(vp, line[0], bold=line[1], italic=line[2], size=line[3])

    # ── THERMAL SPA (row 1, col 2) ──────────────────────────────────────
    spa = main.cell(1, 2)
    box_border(spa)
    set_cell_bg(spa, 'FFFFFF')
    spa.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    nst = spa.add_table(rows=1, cols=1)
    nst.style = 'Table Grid'
    nshc = nst.cell(0, 0)
    set_cell_bg(nshc, '333333')
    box_border(nshc, color='333333')
    nsp = cell_para(nshc, WD_ALIGN_PARAGRAPH.CENTER)
    add_run(nsp, 'THERMAL SPA EXPERIENCE', bold=True, size=9, color='FFFFFF')

    spa_text = ('Discover a dedicated sanctuary for relaxation & rejuvenation, '
                'thoughtfully redesigned with enhanced amenities to elevate your '
                'wellness journey at sea. Unwind all-day with exclusive adults-only '
                "access to Mandara Spa's serene Thermal Suite experience.")
    sp1 = spa.add_paragraph()
    sp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sp1, spa_text, size=8)

    sp2 = spa.add_paragraph()
    sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sp2, 'for only $89/person (day pass)', bold=True, size=11)

    sp3 = spa.add_paragraph()
    sp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sp3, 'from 8am to 10pm | Deck 12 Forward, Spa', italic=True, size=8)

    # ── AIRPORT TRANSFER (row 2, col 0) ─────────────────────────────────
    air = main.cell(2, 0)
    box_border(air)
    set_cell_bg(air, 'FFFFFF')
    air.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    nat = air.add_table(rows=1, cols=1)
    nat.style = 'Table Grid'
    nahc = nat.cell(0, 0)
    set_cell_bg(nahc, '333333')
    box_border(nahc, color='333333')
    nap = cell_para(nahc, WD_ALIGN_PARAGRAPH.CENTER)
    add_run(nap, 'AIRPORT TRANSFER', bold=True, size=9, color='FFFFFF')

    for line in [
        ('From Porto di Ravenna to the following airports:', False, False, 8),
        ('', False, False, 6),
        ('-Venice Marco Polo Airport (VCE)\n-Bologna Airport (BLQ)\n-Forli International Airport (FRL)', False, False, 8),
        ('', False, False, 6),
        ('for only $99.50/person', True, False, 10),
        ('until 10am of May 23 only', False, True, 8),
        ('', False, False, 6),
        ('Visit Guest Services Desk,\nDeck 7 Mid', False, True, 8),
    ]:
        ap = air.add_paragraph()
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(ap, line[0], bold=line[1], italic=line[2], size=line[3])

    # ── TRAIN + LAUNDRY row (row 2, col 2) – split into nested 1x2 ──────
    bot_right = main.cell(2, 2)
    bot_right._tc.clear_content()
    set_cell_border(bot_right, top={}, left={}, bottom={}, right={})

    br_tbl = bot_right.add_table(rows=1, cols=2)
    br_tbl.style = 'Table Grid'

    # Train Station
    tr_cell = br_tbl.cell(0, 0)
    box_border(tr_cell)
    set_cell_bg(tr_cell, 'FFFFFF')
    tr_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    ntr = tr_cell.add_table(rows=1, cols=1)
    ntr.style = 'Table Grid'
    ntrhc = ntr.cell(0, 0)
    set_cell_bg(ntrhc, '333333')
    box_border(ntrhc, color='333333')
    ntrp = cell_para(ntrhc, WD_ALIGN_PARAGRAPH.CENTER)
    add_run(ntrp, 'TRAIN STATION TRANSFER', bold=True, size=8, color='FFFFFF')

    for line in [
        ('From Porto di Ravenna to\nRavenna Train Station', False, False, 8),
        ('for only $23.00/person', True, False, 11),
        ('available until 12nn of May 22', False, False, 8),
        ('Visit Guest Services Desk,\nDeck 7 Mid', False, True, 8),
    ]:
        tp = tr_cell.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(tp, line[0], bold=line[1], italic=line[2], size=line[3])

    # Laundry
    lau_cell = br_tbl.cell(0, 1)
    box_border(lau_cell)
    set_cell_bg(lau_cell, 'FFFFFF')
    lau_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    nlau = lau_cell.add_table(rows=1, cols=1)
    nlau.style = 'Table Grid'
    nlauhc = nlau.cell(0, 0)
    set_cell_bg(nlauhc, '333333')
    box_border(nlauhc, color='333333')
    nlaup = cell_para(nlauhc, WD_ALIGN_PARAGRAPH.CENTER)
    add_run(nlaup, 'LAUNDRY PROMO', bold=True, size=9, color='FFFFFF')

    for line in [
        ('Leave the laundry,\nSeize the day!', True, False, 9),
        ("Fill the bag–we'll handle\nthe rest! Wash & fold in\n48 hours", False, False, 8),
        ('for only $39.50', True, False, 11),
    ]:
        lp = lau_cell.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(lp, line[0], bold=line[1], italic=line[2], size=line[3])

    doc.save('/home/user/sonic-ai-copyx/Onboard_Experiences.docx')
    print('Word saved.')


# ─────────────────────────────────────────────
#  EXCEL FILE
# ─────────────────────────────────────────────
def xl_font(bold=False, size=10, color='000000', italic=False, underline=False):
    return Font(bold=bold, size=size,
                color=color, italic=italic,
                underline='single' if underline else None)


def xl_fill(hex_color):
    return PatternFill('solid', fgColor=hex_color)


def xl_border(style='thin', color='000000'):
    s = Side(style=style, color=color)
    return Border(left=s, right=s, top=s, bottom=s)


def xl_align(h='center', v='center', wrap=True):
    return Alignment(horizontal=h, vertical=v, wrap_text=wrap)


def box(ws, min_row, min_col, max_row, max_col, color='000000', style='medium'):
    s = Side(style=style, color=color)
    n = Side(style=None)
    for r in range(min_row, max_row + 1):
        for c in range(min_col, max_col + 1):
            cell = ws.cell(r, c)
            top = s if r == min_row else n
            bot = s if r == max_row else n
            left = s if c == min_col else n
            right = s if c == max_col else n
            cell.border = Border(top=top, bottom=bot, left=left, right=right)


def header_block(ws, row, col_s, col_e, label):
    """Dark header bar spanning columns."""
    ws.merge_cells(start_row=row, start_column=col_s,
                   end_row=row, end_column=col_e)
    c = ws.cell(row, col_s)
    c.value = label
    c.font = xl_font(bold=True, size=9, color='FFFFFF')
    c.fill = xl_fill('333333')
    c.alignment = xl_align()
    c.border = xl_border()


def write_block(ws, rows_data, start_row, col_s, col_e):
    """Write a list of (text, bold, italic, size, underline) tuples."""
    r = start_row
    for (text, bold, italic, size, underline) in rows_data:
        ws.merge_cells(start_row=r, start_column=col_s,
                       end_row=r, end_column=col_e)
        c = ws.cell(r, col_s)
        c.value = text
        c.font = xl_font(bold=bold, size=size, italic=italic, underline=underline)
        c.alignment = xl_align()
        c.fill = xl_fill('FFFFFF')
        r += 1
    return r


def create_excel():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Onboard Experiences'

    # Column widths (A..N)
    col_widths = {
        'A': 3, 'B': 10, 'C': 10, 'D': 3,
        'E': 10, 'F': 10, 'G': 10, 'H': 10,
        'I': 3, 'J': 10, 'K': 10, 'L': 3,
        'M': 10, 'N': 10,
    }
    for col, w in col_widths.items():
        ws.column_dimensions[col].width = w

    # Row heights
    for i in range(1, 80):
        ws.row_dimensions[i].height = 14

    # ── TITLE ────────────────────────────────────────────────────────────
    ws.merge_cells('B1:G3')
    t = ws['B1']
    t.value = 'ONBOARD EXPERIENCES'
    t.font = xl_font(bold=True, size=22)
    t.alignment = xl_align(h='left', v='center')

    ws.merge_cells('B4:G4')
    ws['B4'].value = 'CORFU, GREECE'
    ws['B4'].font = xl_font(bold=True, size=11)
    ws['B4'].alignment = xl_align(h='left')

    ws.merge_cells('B5:G5')
    ws['B5'].value = 'APPROACH OUR TEAM FOR ASSISTANCE'
    ws['B5'].font = xl_font(bold=True, size=8)
    ws['B5'].alignment = xl_align(h='left')

    # ── FORM FIELDS ──────────────────────────────────────────────────────
    for i, label in enumerate(['GUEST NAME:', 'ROOM NUMBER:',
                                "CREWMEMBER'S NAME & MAPS ID:"], 1):
        r = i * 2 - 1
        ws.merge_cells(start_row=r, start_column=10,
                       end_row=r + 1, end_column=14)
        fc = ws.cell(r, 10)
        fc.value = label
        fc.font = xl_font(bold=True, size=8)
        fc.alignment = xl_align(h='left', v='top')
        box(ws, r, 10, r + 1, 14, style='medium')

    # ── SECTION LAYOUT ───────────────────────────────────────────────────
    # Columns layout:
    # B-D  = left panel (shore / airport)
    # E    = spacer
    # F-N  = right panel (vibe / spa / train / laundry)

    # ── SHORE EXCURSION ──────────────────────────────────────────────────
    SH_R1 = 7
    SH_CS, SH_CE = 2, 4  # cols B-D

    header_block(ws, SH_R1, SH_CS, SH_CE, 'SHORE EXCURSION')

    shore_data = [
        ('Kotor Tours', False, False, 9, True),
        ('BOKA BAY CATAMARAN SAIL & BEACH', True, False, 8, False),
        ('$199.99Adult | $149.99 Child', False, False, 8, False),
        ('HISTORIC KOTOR & BUDVA', True, False, 8, False),
        ('$99.99Adult | $79.99 Child', False, False, 8, False),
        ('', False, False, 8, False),
        ('Split Tours', False, False, 9, True),
        ('TROGIR & BEACH', True, False, 8, False),
        ('$89.99Adult | $49.99 Child', False, False, 8, False),
        ('SPLIT OPEN-AIR SIGHTSEEING & WALKING TOUR', True, False, 8, False),
        ('$69.99Adult | $49.99 Child', False, False, 8, False),
        ('', False, False, 8, False),
        ('Visit Shore Excursions Desk Deck 7, Mid', False, True, 8, False),
    ]
    shore_end = write_block(ws, shore_data, SH_R1 + 1, SH_CS, SH_CE)
    box(ws, SH_R1, SH_CS, shore_end - 1, SH_CE)

    # ── VIBE BEACH CLUB ──────────────────────────────────────────────────
    VB_R1 = 7
    VB_CS, VB_CE = 6, 14

    header_block(ws, VB_R1, VB_CS, VB_CE, 'VIBE BEACH CLUB')

    vibe_data = [
        ('Escape to luxury at sea.', False, False, 9, False),
        ('Unwind in style with an exclusive, adults-only retreat with ocean views & enjoy premium touches designed just for you.', False, False, 8, False),
        ('Limited access available.', False, False, 8, False),
        ('Reserve early for the ultimate experience.', False, False, 8, False),
        ('Vibe Cruise Pass - $209/person (4 days)', True, False, 9, False),
        ('Vibe Cabanas - $499/2pax (4 days)', True, False, 9, False),
        ('', False, False, 8, False),
        ('Visit Guest Service Desk, Deck 7 Midship', False, True, 8, False),
    ]
    vibe_end = write_block(ws, vibe_data, VB_R1 + 1, VB_CS, VB_CE)
    box(ws, VB_R1, VB_CS, vibe_end - 1, VB_CE)

    # ── THERMAL SPA ──────────────────────────────────────────────────────
    SPA_R1 = vibe_end + 1
    SPA_CS, SPA_CE = 6, 14

    header_block(ws, SPA_R1, SPA_CS, SPA_CE, 'THERMAL SPA EXPERIENCE')

    spa_data = [
        ('Discover a dedicated sanctuary for relaxation & rejuvenation, thoughtfully redesigned with enhanced amenities to elevate your wellness journey at sea. Unwind all-day with exclusive adults-only access to Mandara Spa\'s serene Thermal Suite experience.', False, False, 8, False),
        ('for only $89/person (day pass)', True, False, 12, False),
        ('from 8am to 10pm | Deck 12 Forward, Spa', False, True, 8, False),
    ]
    spa_end = write_block(ws, spa_data, SPA_R1 + 1, SPA_CS, SPA_CE)
    box(ws, SPA_R1, SPA_CS, spa_end - 1, SPA_CE)

    # Adjust Spa rows to be taller
    for r in range(SPA_R1, spa_end):
        ws.row_dimensions[r].height = 28

    # ── AIRPORT TRANSFER ─────────────────────────────────────────────────
    AIR_R1 = shore_end + 1
    AIR_CS, AIR_CE = 2, 4

    header_block(ws, AIR_R1, AIR_CS, AIR_CE, 'AIRPORT TRANSFER')

    air_data = [
        ('From Porto di Ravenna to the following airports:', False, False, 8, False),
        ('', False, False, 6, False),
        ('-Venice Marco Polo Airport (VCE)', False, False, 8, False),
        ('-Bologna Airport (BLQ)', False, False, 8, False),
        ('-Forli International Airport (FRL)', False, False, 8, False),
        ('', False, False, 6, False),
        ('for only $99.50/person', True, False, 11, False),
        ('until 10am of May 23 only', False, True, 8, False),
        ('', False, False, 6, False),
        ('Visit Guest Services Desk, Deck 7 Mid', False, True, 8, False),
    ]
    air_end = write_block(ws, air_data, AIR_R1 + 1, AIR_CS, AIR_CE)
    box(ws, AIR_R1, AIR_CS, air_end - 1, AIR_CE)

    # ── TRAIN STATION TRANSFER ───────────────────────────────────────────
    TR_R1 = spa_end + 1
    TR_CS, TR_CE = 6, 9

    header_block(ws, TR_R1, TR_CS, TR_CE, 'TRAIN STATION TRANSFER')

    tr_data = [
        ('From Porto di Ravenna to Ravenna Train Station', False, False, 8, False),
        ('for only $23.00/person', True, False, 11, False),
        ('available until 12nn of May 22', False, False, 8, False),
        ('Visit Guest Services Desk, Deck 7 Mid', False, True, 8, False),
    ]
    tr_end = write_block(ws, tr_data, TR_R1 + 1, TR_CS, TR_CE)
    box(ws, TR_R1, TR_CS, tr_end - 1, TR_CE)

    # ── LAUNDRY PROMO ────────────────────────────────────────────────────
    LAU_R1 = spa_end + 1
    LAU_CS, LAU_CE = 11, 14

    header_block(ws, LAU_R1, LAU_CS, LAU_CE, 'LAUNDRY PROMO')

    lau_data = [
        ('Leave the laundry, Seize the day!', True, False, 9, False),
        ("Fill the bag–we'll handle the rest! Wash & fold in 48 hours", False, False, 8, False),
        ('for only $39.50', True, False, 11, False),
    ]
    lau_end = write_block(ws, lau_data, LAU_R1 + 1, LAU_CS, LAU_CE)
    box(ws, LAU_R1, LAU_CS, lau_end - 1, LAU_CE)

    # Taller rows for bottom section
    for r in range(TR_R1, max(tr_end, lau_end)):
        ws.row_dimensions[r].height = 22

    # Freeze top rows
    ws.freeze_panes = 'B7'

    wb.save('/home/user/sonic-ai-copyx/Onboard_Experiences.xlsx')
    print('Excel saved.')


create_word()
create_excel()
print('Done!')
